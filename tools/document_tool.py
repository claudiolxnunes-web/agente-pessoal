"""Análise de documentos: PDF, DOCX, XLSX, CSV, TXT"""
import os
import csv
import io
from typing import List, Dict, Any, Optional
from pathlib import Path

class DocumentTool:
    """Ferramenta para analisar e extrair informações de documentos"""

    def __init__(self):
        self.documentos_carregados = {}

    def analisar_pdf(self, caminho: str) -> str:
        """Extrai texto de PDF"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(caminho)
            texto = ""

            for i, page in enumerate(reader.pages):
                texto += f"\n--- Página {i+1} ---\n"
                texto += page.extract_text()

            self.documentos_carregados[caminho] = {"tipo": "pdf", "conteudo": texto}

            resumo = texto[:2000] + "..." if len(texto) > 2000 else texto
            return f"📄 PDF analisado: {len(reader.pages)} páginas\n\n{resumo}"

        except ImportError:
            return "❌ Instale pypdf: pip install pypdf"
        except Exception as e:
            return f"Erro ao ler PDF: {e}"

    def analisar_docx(self, caminho: str) -> str:
        """Extrai texto de documento Word"""
        try:
            from docx import Document

            doc = Document(caminho)
            texto = []

            for para in doc.paragraphs:
                if para.text.strip():
                    texto.append(para.text)

            # Tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    texto.append(row_text)

            conteudo = "\n".join(texto)
            self.documentos_carregados[caminho] = {"tipo": "docx", "conteudo": conteudo}

            resumo = conteudo[:2000] + "..." if len(conteudo) > 2000 else conteudo
            return f"📝 DOCX analisado\n\n{resumo}"

        except ImportError:
            return "❌ Instale python-docx: pip install python-docx"
        except Exception as e:
            return f"Erro ao ler DOCX: {e}"

    def analisar_csv(self, caminho: str) -> str:
        """Analisa arquivo CSV e retorna estatísticas"""
        try:
            with open(caminho, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                headers = next(reader)

                dados = []
                for row in reader:
                    dados.append(row)

                # Estatísticas básicas
                num_linhas = len(dados)
                num_colunas = len(headers)

                # Análise de colunas numéricas
                analise = []
                for i, header in enumerate(headers):
                    valores = [row[i] for row in dados if i < len(row)]

                    # Tenta converter para número
                    try:
                        nums = [float(v.replace(',', '.')) for v in valores if v]
                        if nums:
                            media = sum(nums) / len(nums)
                            analise.append(f"📊 {header}: média={media:.2f}, min={min(nums):.2f}, max={max(nums):.2f}")
                    except:
                        # Coluna categórica
                        unicos = set(valores)
                        analise.append(f"📋 {header}: {len(unicos)} valores únicos")

                self.documentos_carregados[caminho] = {
                    "tipo": "csv", 
                    "headers": headers,
                    "dados": dados[:100]  # Primeiras 100 linhas
                }

                output = f"📊 CSV analisado: {num_linhas} linhas × {num_colunas} colunas\n\n"
                output += "Colunas:\n" + "\n".join([f"• {h}" for h in headers]) + "\n\n"
                output += "Análise:\n" + "\n".join(analise[:10])

                return output

        except Exception as e:
            return f"Erro ao ler CSV: {e}"

    def analisar_xlsx(self, caminho: str) -> str:
        """Analisa planilha Excel"""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(caminho, data_only=True)

            output = f"📈 Excel analisado: {len(wb.sheetnames)} abas\n\n"

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                output += f"📑 Aba: {sheet_name} ({ws.max_row} linhas × {ws.max_column} cols)\n"

                # Primeiras linhas como preview
                preview = []
                for row in ws.iter_rows(min_row=1, max_row=min(5, ws.max_row), values_only=True):
                    preview.append(" | ".join([str(v) if v else "" for v in row]))

                output += "\n".join(preview) + "\n\n"

            self.documentos_carregados[caminho] = {"tipo": "xlsx", "workbook": wb}

            return output

        except ImportError:
            return "❌ Instale openpyxl: pip install openpyxl"
        except Exception as e:
            return f"Erro ao ler Excel: {e}"

    def analisar_txt(self, caminho: str) -> str:
        """Lê arquivo de texto"""
        try:
            with open(caminho, 'r', encoding='utf-8') as f:
                conteudo = f.read()

            self.documentos_carregados[caminho] = {"tipo": "txt", "conteudo": conteudo}

            linhas = conteudo.split('\n')
            resumo = conteudo[:2000] + "..." if len(conteudo) > 2000 else conteudo

            return f"📄 TXT: {len(linhas)} linhas, {len(conteudo)} caracteres\n\n{resumo}"

        except Exception as e:
            return f"Erro ao ler TXT: {e}"

    def analisar_auto(self, caminho: str) -> str:
        """Detecta tipo e analisa automaticamente"""
        ext = Path(caminho).suffix.lower()

        analisadores = {
            '.pdf': self.analisar_pdf,
            '.docx': self.analisar_docx,
            '.csv': self.analisar_csv,
            '.xlsx': self.analisar_xlsx,
            '.xls': self.analisar_xlsx,
            '.txt': self.analisar_txt,
            '.md': self.analisar_txt,
        }

        if ext in analisadores:
            return analisadores[ext](caminho)
        else:
            return f"❌ Formato '{ext}' não suportado. Use: PDF, DOCX, CSV, XLSX, TXT"

    def perguntar_sobre_documento(self, caminho: str, pergunta: str, llm) -> str:
        """Faz perguntas sobre um documento carregado usando LLM"""
        if caminho not in self.documentos_carregados:
            return "Documento não carregado. Analise-o primeiro."

        doc = self.documentos_carregados[caminho]
        conteudo = doc.get("conteudo", "")

        if not conteudo:
            return "Não foi possível extrair conteúdo do documento."

        # Limita contexto para não exceder tokens
        contexto = conteudo[:4000]

        prompt = f"""Com base no documento abaixo, responda à pergunta.

DOCUMENTO:
{contexto}

PERGUNTA: {pergunta}

Responda de forma concisa e direta."""

        resposta = llm.invoke(prompt)
        return resposta.content

    def listar_documentos(self) -> str:
        """Lista documentos carregados"""
        if not self.documentos_carregados:
            return "Nenhum documento carregado."

        output = "📚 Documentos carregados:\n"
        for caminho, info in self.documentos_carregados.items():
            nome = Path(caminho).name
            tipo = info["tipo"].upper()
            output += f"• {nome} ({tipo})\n"

        return output


# ── Função de módulo ──────────────────────────────────────────────────────────
def analisar_documento(caminho: str, pergunta: str = "") -> str:
    try:
        resultado = DocumentTool().analisar_auto(caminho)
        return resultado
    except Exception as e:
        return f"Erro ao analisar documento: {e}"
